from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DEFAULT_FONT = "Microsoft JhengHei"
HEADING_COLOR = RGBColor(31, 78, 121)


def set_run_font(run, size: float | None = None, bold: bool = False, italic: bool = False) -> None:
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), DEFAULT_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), DEFAULT_FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), DEFAULT_FONT)
    run.font.bold = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)


def set_paragraph_font(paragraph, size: float | None = None, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def configure_document(document: Document) -> None:
    for style_name in [
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Bullet",
    ]:
        style = document.styles[style_name]
        style.font.name = DEFAULT_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), DEFAULT_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DEFAULT_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), DEFAULT_FONT)

    document.styles["Normal"].font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True)
    run.font.color.rgb = HEADING_COLOR


def add_quote(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, italic=True)
    run.font.color.rgb = RGBColor(68, 114, 196)


def parse_markdown_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start_index

    while index < len(lines):
        stripped = lines[index].strip()
        if not (stripped.startswith("|") and stripped.endswith("|")):
            break

        cells = [cell.strip().strip("`") for cell in stripped.strip("|").split("|")]
        is_alignment_row = all(set(cell) <= {"-", ":"} for cell in cells)
        if not is_alignment_row:
            rows.append(cells)
        index += 1

    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            text = row[column_index] if column_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            set_run_font(run, size=10, bold=row_index == 0)


def add_markdown_lines(document: Document, lines: list[str]) -> None:
    in_code_block = False
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if in_code_block:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(stripped)
            set_run_font(run, size=9)
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, next_index = parse_markdown_table(lines, index)
            add_table(document, rows)
            index = next_index
            continue

        if stripped.startswith("# "):
            add_heading(document, stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            add_heading(document, stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(stripped[2:].strip())
            set_run_font(run, size=10.5)
        elif stripped.startswith(">"):
            quote = stripped.lstrip("> ").strip()
            if quote:
                add_quote(document, quote)
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(stripped)
            set_run_font(run, size=10.5)

        index += 1


def export_docx(source: Path, output: Path) -> Path:
    if not source.exists():
        raise FileNotFoundError(f"Source markdown not found: {source}")

    document = Document()
    configure_document(document)

    lines = source.read_text(encoding="utf-8-sig").splitlines()
    add_markdown_lines(document, lines)

    output.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(output)
    except PermissionError as exc:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        fallback = output.with_name(f"{output.stem}-{timestamp}{output.suffix}")
        document.save(fallback)
        print(
            f"Cannot write {output} because it is locked. "
            f"Exported fallback Word report: {fallback}"
        )
        return fallback

    return output


def main() -> None:
    parser = argparse.ArgumentParser(description="Export PM markdown report to Word docx.")
    parser.add_argument(
        "--source",
        default="artifacts/generated/pm/release-summary.md",
        help="Source PM markdown report.",
    )
    parser.add_argument(
        "--output",
        default="artifacts/generated/pm/release-summary.docx",
        help="Output Word docx path.",
    )
    args = parser.parse_args()

    actual_output = export_docx(Path(args.source), Path(args.output))
    print(f"Exported PM Word report: {actual_output}")


if __name__ == "__main__":
    main()
